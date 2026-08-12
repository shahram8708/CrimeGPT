from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.models.evidence import CaseDiaryEntry
from app.services.file_service import generated_files
from app.utils.formatting_utils import format_ist

NOTICE = (
    "Draft assistance extract. The officer remains responsible. "
    "AI-generated legal information may contain errors and should be verified "
    "against authoritative legal sources. This platform does not provide legal advice."
)

HEADINGS = {
    "en": {
        "title": "Case diary extract",
        "station": "Station",
        "cr": "CR / FIR",
        "io": "Investigating officer",
        "range": "Period",
        "type": "Type",
        "when": "Occurred",
        "author": "Author",
        "place": "Place",
        "status": "Status",
        "corrects": "Corrects entry",
    },
    "hi": {
        "title": "केस डायरी उद्धरण",
        "station": "थाना",
        "cr": "सीआर / एफआईआर",
        "io": "अन्वेषण अधिकारी",
        "range": "अवधि",
        "type": "प्रकार",
        "when": "समय",
        "author": "लेखक",
        "place": "स्थान",
        "status": "स्थिति",
        "corrects": "सुधार प्रविष्टि",
    },
    "gu": {
        "title": "કેસ ડાયરી ઉતારો",
        "station": "થાણું",
        "cr": "સીઆર / એફઆઇઆર",
        "io": "તપાસ અધિકારી",
        "range": "ગાળો",
        "type": "પ્રકાર",
        "when": "સમય",
        "author": "લેખક",
        "place": "સ્થળ",
        "status": "સ્થિતિ",
        "corrects": "સુધારે છે",
    },
}


def _entries(exp):
    from datetime import datetime, time, timezone

    rows = CaseDiaryEntry.query.filter(CaseDiaryEntry.case_id == exp.case_id).all()
    start = datetime.combine(exp.date_from, time.min, tzinfo=timezone.utc)
    end = datetime.combine(exp.date_to, time.max, tzinfo=timezone.utc)
    picked = []
    for e in rows:
        at = e.occurred_at
        if at.tzinfo is None:
            at = at.replace(tzinfo=timezone.utc)
        if start <= at <= end:
            picked.append(e)
            if e.corrects_entry_id and e.corrects and e.corrects not in picked:
                picked.append(e.corrects)
    picked.sort(key=lambda e: e.occurred_at)
    return picked


def ensure_diary_template():
    folder = Path(__file__).resolve().parent.parent / "templates_docx"
    folder.mkdir(parents=True, exist_ok=True)
    path = folder / "diary_extract_en.docx"
    if path.is_file() and path.stat().st_size > 1000:
        return path
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("{{station_name}}")
    run.bold = True
    doc.add_paragraph("{{letterhead_2}}")
    doc.add_paragraph("{{letterhead_3}}")
    doc.add_heading("Case diary extract", level=1)
    doc.add_paragraph("{{notice}}")
    doc.save(path)
    return path


def render_diary_docx(exp):
    labels = HEADINGS.get(exp.language) or HEADINGS["en"]
    case = exp.case
    station = case.station
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    head = doc.add_paragraph()
    r = head.add_run(station.name if station else "CrimeGPT")
    r.bold = True
    r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)
    if station and station.letterhead_line2:
        doc.add_paragraph(station.letterhead_line2)
    if station and station.letterhead_line3:
        doc.add_paragraph(station.letterhead_line3)
    title = doc.add_paragraph()
    tr = title.add_run(labels["title"])
    tr.bold = True
    tr.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    notice = doc.add_paragraph(NOTICE)
    notice.runs[0].italic = True
    io = case.assigned_io.full_name if case.assigned_io else "-"
    meta = (
        f"{labels['station']}: {station.name if station else '-'} ({station.code if station else ''})\n"
        f"{labels['cr']}: {case.display_cr}\n"
        f"{labels['io']}: {io}\n"
        f"{labels['range']}: {exp.date_from.isoformat()} – {exp.date_to.isoformat()}"
    )
    doc.add_paragraph(meta)
    if getattr(exp, "summary_text", None):
        sh = doc.add_paragraph()
        run = sh.add_run("AI summary - verify")
        run.bold = True
        doc.add_paragraph(NOTICE)
        doc.add_paragraph(exp.summary_text)
    for entry in _entries(exp):
        doc.add_paragraph("-" * 24)
        block = (
            f"{labels['type']}: {entry.entry_type}    {labels['status']}: {entry.status}\n"
            f"{labels['when']}: {format_ist(entry.occurred_at)}    "
            f"{labels['author']}: {entry.author.full_name if entry.author else '-'}\n"
            f"{labels['place']}: {entry.place or '-'}"
        )
        if entry.corrects_entry_id:
            block += f"\n{labels['corrects']}: {entry.corrects_entry_id}"
        doc.add_paragraph(block)
        doc.add_paragraph(entry.body or "")
    key = f"diary/{exp.uuid}.docx"
    generated_files.put(b"", key)
    path = generated_files._resolve(key)
    doc.save(path)
    return key


def _templates_dir():
    folder = Path(__file__).resolve().parent.parent / "templates_docx"
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def _write_simple_template(path, title):
    doc = Document()
    doc.add_paragraph("{{station_name}}")
    doc.add_paragraph("{{letterhead_2}}")
    doc.add_paragraph("{{letterhead_3}}")
    doc.add_heading(title, level=1)
    doc.add_paragraph("{{disclaimer}}")
    doc.add_paragraph("{{notice}}")
    doc.add_paragraph("{{body_block}}")
    doc.save(path)


def ensure_court_templates():
    folder = _templates_dir()
    titles = {
        "medical_letter": "Medical Treatment Letter",
        "seizure_receipt": "Seizure Receipt",
        "remand_pc": "Remand Request - Police Custody",
        "face_identification": "Accused Face Identification Form",
        "purvani_chargesheet": "Purvani Chargesheet",
        "court_custody": "Court Custody Letter",
        "accused_panchanama": "Accused Panchanama",
        "lers_request": "Legal request letter (template)",
    }
    for key, title in titles.items():
        for lang in ("en", "hi", "gu"):
            path = folder / f"{key}_{lang}.docx"
            if not path.is_file() or path.stat().st_size < 800:
                _write_simple_template(path, title)
    return folder


def _body_block(ctx):
    kind = ctx.get("doc_type")
    lines = []
    if kind == "medical_letter":
        inj = ctx.get("injured") or {}
        ai = ctx.get("ai") or {}
        lines = [
            f"To: The Medical Officer, {ctx.get('hospital_name')}",
            f"Injured / complainant: {inj.get('name')}  Age: {inj.get('age')}  Address: {inj.get('address')}",
            f"Incident: {ctx.get('incident_at')} at {ctx.get('place')}",
            f"Alleged history: {ai.get('history_line') or ctx.get('history_reported') or ctx.get('gist')}",
            f"Requested examination: {ai.get('request_line') or ctx.get('requested_exam') or 'Wound certificate'}",
            f"Escorting officer: {ctx.get('escort_name')}",
            f"Station: {ctx.get('station_name')} ({ctx.get('station_code')})",
            f"Requesting officer: {ctx.get('officer_name')}, {ctx.get('officer_rank')}",
            f"Investigating officer: {ctx.get('io_name')}",
        ]
    elif kind == "seizure_receipt":
        lines = [
            f"Place of seizure: {ctx.get('place')}    CR / FIR: {ctx.get('cr')}",
            "Articles seized:",
        ]
        for item in ctx.get("items") or []:
            lines.append(
                f"- {item.get('description')}  Qty {item.get('quantity')} {item.get('unit')}  "
                f"Marking {item.get('marking')}  Exhibit {item.get('exhibit')}"
            )
        accused = ctx.get("accused") or []
        if accused:
            lines.append("Accused / possessor: " + ", ".join(a.get("name") for a in accused))
        panch = ctx.get("panch") or []
        for i, p in enumerate(panch[:2], 1):
            lines.append(f"Panch {i}: {p.get('name')}, {p.get('address')}")
        lines.append(f"Seizing officer: {ctx.get('officer_name')}, {ctx.get('officer_rank')}")
        lines.append("Signatures: Seizing officer ________    Panch 1 ________    Panch 2 ________")
    elif kind == "remand_pc":
        ai = ctx.get("ai") or {}
        lines = [
            f"To: {ctx.get('court_name') or 'The learned Magistrate'}",
            f"CR / FIR: {ctx.get('cr')}    Station: {ctx.get('station_name')}",
        ]
        for a in ctx.get("accused") or []:
            lines.append(f"Accused: {a.get('name')} s/o {a.get('guardian')}, age {a.get('age')}, {a.get('address')}")
        for ar in ctx.get("arrests") or []:
            lines.append(f"Arrest: {ar.get('accused_name')} at {ar.get('arrest_at')}, {ar.get('place')}")
        secs = ctx.get("sections") or []
        if secs:
            lines.append("Confirmed sections: " + ", ".join(f"{s.get('family')} {s.get('code')}" for s in secs))
        grounds = ai.get("grounds_of_custody") or []
        if isinstance(grounds, list) and grounds:
            lines.append("Grounds of custody:")
            lines.extend(f"- {g}" for g in grounds)
        elif ctx.get("grounds"):
            lines.append("Grounds: " + ctx.get("grounds"))
        pending = ai.get("investigation_pending") or []
        if pending:
            lines.append("Investigation pending:")
            lines.extend(f"- {p}" for p in pending)
        prayer = ai.get("prayer") or f"Police custody for {ctx.get('custody_hours') or '-'} hours is sought."
        lines.append("Prayer: " + prayer)
        lines.append("I verify that the facts above are true to my knowledge.")
        lines.append(f"IO: {ctx.get('io_name')}    {ctx.get('generated_at')}")
    elif kind == "face_identification":
        ai = ctx.get("ai") or {}
        lines = [
            f"Case: {ctx.get('cr')}    Station: {ctx.get('station_name')}",
            ctx.get("photo_note") or "Photograph to be affixed in the box provided.",
        ]
        for a in ctx.get("accused") or []:
            lines.append(
                f"Name: {a.get('name')}  Alias: {a.get('alias')}  Guardian: {a.get('guardian')}  "
                f"Age: {a.get('age')}  Gender: {a.get('gender')}"
            )
        if ai.get("identifiers"):
            lines.append("Identifiers: " + str(ai.get("identifiers")))
        lines.append(f"Investigating officer: {ctx.get('io_name')}    Date: {ctx.get('generated_at')}")
    elif kind == "purvani_chargesheet":
        ai = ctx.get("ai") or {}
        lines = [
            f"Purvani / chargesheet outline    CR / FIR: {ctx.get('cr')}",
            f"Station: {ctx.get('station_name')} ({ctx.get('station_code')})",
            f"Complainant: {', '.join(p.get('name') for p in (ctx.get('complainants') or [])) or '-'}",
            f"Narrative gist: {ai.get('gist') or ctx.get('gist')}",
        ]
        for a in ctx.get("accused") or []:
            lines.append(f"Accused: {a.get('name')} s/o {a.get('guardian')}, {a.get('address')}")
        secs = ctx.get("sections") or []
        if secs:
            lines.append("Confirmed sections: " + ", ".join(f"{s.get('family')} {s.get('code')}" for s in secs))
        wits = ctx.get("witnesses") or []
        if wits:
            lines.append("Witnesses: " + ", ".join(w.get("name") for w in wits))
        for item in ctx.get("items") or []:
            lines.append(f"Exhibit {item.get('exhibit')}: {item.get('description')}")
        lines.append("Investigation: " + (ai.get("investigation_outline") or ctx.get("investigation_gist") or "-"))
        lines.append("Forwarding: " + (ai.get("prayer") or "Submitted for further process."))
        lines.append(f"IO: {ctx.get('io_name')}")
    elif kind == "court_custody":
        lines = [
            f"To: {ctx.get('court_name') or 'The learned Magistrate'}",
            f"Production of accused in judicial custody    CR / FIR: {ctx.get('cr')}",
            f"Production: {ctx.get('production_at') or 'as directed'}",
        ]
        for a in ctx.get("accused") or []:
            lines.append(f"Accused: {a.get('name')} s/o {a.get('guardian')}, age {a.get('age')}")
        secs = ctx.get("sections") or []
        if secs:
            lines.append("Confirmed sections: " + ", ".join(f"{s.get('family')} {s.get('code')}" for s in secs))
        lines.append("Articles accompanying:")
        for item in ctx.get("items") or []:
            lines.append(f"- {item.get('description')} ({item.get('exhibit')})")
        if not ctx.get("items"):
            lines.append("- Case papers and the accused as produced.")
        lines.append("Request: The accused may be taken in court custody.")
        lines.append(f"Escorting officer: {ctx.get('officer_name')}")
        lines.append(f"IO: {ctx.get('io_name')}")
    elif kind == "accused_panchanama":
        lines = [
            f"Accused panchanama    Place: {ctx.get('place')}    Time: {ctx.get('incident_at')}",
            f"CR / FIR: {ctx.get('cr')}",
        ]
        for a in ctx.get("accused") or []:
            lines.append(
                f"Accused: {a.get('name')} s/o {a.get('guardian')}, age {a.get('age')}, "
                f"gender {a.get('gender')}, {a.get('address')}"
            )
        desc = (ctx.get("ai") or {}).get("description") or ctx.get("person_description") or "As observed at the place."
        lines.append("Description of person, clothes and visible injuries: " + desc)
        for i, p in enumerate((ctx.get("panch") or [])[:2], 1):
            lines.append(f"Panch {i}: {p.get('name')}, {p.get('address')}")
        lines.append(f"IO: {ctx.get('io_name')}    {ctx.get('generated_at')}")
        lines.append("Signatures: IO ________    Panch 1 ________    Panch 2 ________")
    elif kind == "lers_request":
        ai = ctx.get("ai") or {}
        lines = [
            "Legal request letter (template)",
            "This letter is a station draft. It has not been filed with any platform.",
            f"To: {ctx.get('request_target') or 'the service provider'}",
            f"From: {ctx.get('station_name')} ({ctx.get('station_code')})",
            f"CR / FIR: {ctx.get('cr')}    Period: {ctx.get('date_from') or '-'} to {ctx.get('date_to') or '-'}",
        ]
        for a in ctx.get("accused") or []:
            lines.append(f"Subject person already on the case: {a.get('name')} s/o {a.get('guardian')}")
        secs = ctx.get("sections") or []
        if secs:
            lines.append("Statutory hook: " + ", ".join(f"{s.get('family')} {s.get('code')}" for s in secs))
        lines.append(
            ai.get("request_paragraph")
            or "Please preserve and disclose account or device records for the period above, for use in this investigation."
        )
        lines.append("No account identifier is invented. Only identifiers already typed on the case are used.")
        lines.append(f"Requesting IO: {ctx.get('io_name')}")
    if ctx.get("is_incomplete"):
        lines.insert(0, "DRAFT-INCOMPLETE - missing required pool fields.")
    return "\n".join(lines)


def render_court_docx(doc_row, ctx):
    ensure_court_templates()
    folder = _templates_dir()
    lang = ctx.get("language") or "en"
    kind = ctx.get("doc_type")
    path = folder / f"{kind}_{lang}.docx"
    if not path.is_file():
        path = folder / f"{kind}_en.docx"
    if not path.is_file():
        raise RuntimeError(f"Template missing for {kind}")
    from docxtpl import DocxTemplate

    payload = dict(ctx)
    payload["body_block"] = _body_block(ctx)
    tpl = DocxTemplate(str(path))
    try:
        tpl.render(payload)
    except Exception:
        doc = Document(str(path))
        doc.add_paragraph(payload["body_block"])
        key = f"docs/{doc_row.uuid}.docx"
        generated_files.put(b"", key)
        doc.save(generated_files._resolve(key))
        return key
    key = f"docs/{doc_row.uuid}.docx"
    generated_files.put(b"", key)
    tpl.save(generated_files._resolve(key))
    return key
